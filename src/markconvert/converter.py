"""Document conversion using Ollama LLM for PDFs/images and Docling for office formats."""

import io
import os
import tempfile
from pathlib import Path
from typing import Union, BinaryIO, Optional

from docling.document_converter import DocumentConverter
from docx import Document
from docx.shared import Pt
import markdown
from weasyprint import HTML

from markconvert.ollama_client import OllamaClient
from markconvert.pdf_to_image import PdfToImageConverter, get_pdf_page_count


class MarkdownConverter:
    """Handle conversion between Markdown and other document formats."""

    # File extensions that are processed via LLM (vision)
    LLM_FORMATS = {'.pdf', '.png', '.jpg', '.jpeg'}

    # File extensions that are processed via Docling
    DOCLING_FORMATS = {'.docx', '.doc', '.pptx', '.ppt', '.html', '.htm'}

    # Text formats that don't need processing
    TEXT_FORMATS = {'.txt', '.md'}

    def __init__(
        self,
        ollama_url: str = "http://localhost:11434",
        ollama_model: str = "gemma3:27b",
        ollama_vision_model: str = "qwen2.5vl:32b"
    ):
        """
        Initialize the converter.

        Args:
            ollama_url: Ollama server URL
            ollama_model: Default text model
            ollama_vision_model: Vision model for image/PDF processing
        """
        # Initialize Ollama client for PDF/image processing
        self.ollama_client = OllamaClient(
            base_url=ollama_url,
            model=ollama_model,
            vision_model=ollama_vision_model
        )

        # Initialize Docling for office document processing
        self.doc_converter = DocumentConverter()

        # Initialize PDF-to-image converter with optimized settings
        self.pdf_converter = PdfToImageConverter(
            dpi=150,  # Lower DPI = faster processing
            output_format="jpeg",  # JPEG uses fewer tokens than PNG
            jpeg_quality=85,  # Good quality, smaller file
            max_dimension=2048  # Limit image size for faster processing
        )

    def import_document(self, file_path: str) -> str:
        """
        Import a document and convert to Markdown.

        Routing logic:
        - PDF/PNG/JPG → Ollama Vision LLM
        - DOCX/PPTX/HTML → Docling
        - TXT/MD → Direct read

        Args:
            file_path: Path to the document file

        Returns:
            Markdown text
        """
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        try:
            # Route 1: PDF and images → Ollama Vision
            if suffix in self.LLM_FORMATS:
                return self._import_via_llm(file_path)

            # Route 2: Office formats → Docling
            elif suffix in self.DOCLING_FORMATS:
                return self._import_via_docling(file_path)

            # Route 3: Plain text formats → Direct read
            elif suffix in self.TEXT_FORMATS:
                return file_path.read_text(encoding='utf-8')

            else:
                raise ValueError(
                    f"Unsupported file format: {suffix}. "
                    f"Supported: {', '.join(sorted(self.LLM_FORMATS | self.DOCLING_FORMATS | self.TEXT_FORMATS))}"
                )

        except Exception as e:
            raise ValueError(f"Fehler beim Importieren der Datei: {str(e)}")

    def _import_via_llm(self, file_path: Path) -> str:
        """Import PDF or image via Ollama Vision LLM."""
        suffix = file_path.suffix.lower()

        if suffix == '.pdf':
            # PDF: Convert to images, process each page
            return self._process_pdf_via_llm(file_path)
        else:
            # Image: Process directly
            return self.ollama_client.process_image_to_markdown(file_path)

    def _process_pdf_via_llm(self, pdf_path: Path) -> str:
        """Process PDF by converting pages to images and analyzing with LLM."""
        import time
        import logging

        logger = logging.getLogger(__name__)

        # Get page count
        page_count = get_pdf_page_count(pdf_path)
        logger.info(f"Starting PDF processing: {page_count} pages")

        # Create temp directory for images
        temp_dir = Path(tempfile.mkdtemp(prefix="pdf_pages_"))

        try:
            # Convert all pages to images
            start_time = time.time()
            logger.info("Converting PDF to images...")
            image_paths = self.pdf_converter.convert_pdf_to_images(pdf_path, temp_dir)
            conversion_time = time.time() - start_time
            logger.info(f"PDF converted to images in {conversion_time:.2f}s")

            # Log image sizes for debugging
            total_size = sum(p.stat().st_size for p in image_paths)
            logger.info(f"Total image size: {total_size / 1024 / 1024:.2f} MB ({total_size / len(image_paths) / 1024:.2f} KB/page)")

            # Process each page
            page_markdowns = []
            for i, image_path in enumerate(image_paths, start=1):
                page_start = time.time()
                logger.info(f"Processing page {i}/{page_count}...")
                print(f"Processing page {i}/{page_count}...")

                # Process image to markdown
                page_md = self.ollama_client.process_image_to_markdown(image_path)

                page_time = time.time() - page_start
                logger.info(f"Page {i} processed in {page_time:.2f}s")

                # Add page separator
                if i > 1:
                    page_markdowns.append(f"\n\n---\n\n## Page {i}\n\n")
                page_markdowns.append(page_md)

                # Sleep between pages to give GPU time to recover (except after last page)
                if i < page_count:
                    logger.info("Waiting 10 seconds before next page...")
                    time.sleep(10)

            total_time = time.time() - start_time
            logger.info(f"Total processing time: {total_time:.2f}s ({total_time/page_count:.2f}s/page)")

            # Combine all pages
            return "".join(page_markdowns)

        finally:
            # Clean up temp images
            import shutil
            shutil.rmtree(temp_dir, ignore_errors=True)

    def _import_via_docling(self, file_path: Path) -> str:
        """Import office document via Docling."""
        result = self.doc_converter.convert(str(file_path))
        return result.document.export_to_markdown()

    def import_document_from_bytes(self, file_bytes: bytes, filename: str) -> str:
        """
        Import a document from bytes and convert to Markdown.

        Args:
            file_bytes: Document content as bytes
            filename: Original filename (for format detection)

        Returns:
            Markdown text
        """
        # Save to temporary file
        suffix = Path(filename).suffix
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            markdown_text = self.import_document(tmp_path)
            return markdown_text
        finally:
            # Clean up temporary file
            Path(tmp_path).unlink(missing_ok=True)

    def export_to_docx(self, markdown_text: str) -> bytes:
        """
        Export Markdown to DOCX format.

        Args:
            markdown_text: Markdown content

        Returns:
            DOCX file as bytes
        """
        doc = Document()

        # Parse markdown line by line
        lines = markdown_text.split('\n')

        for line in lines:
            line = line.rstrip()

            # Headers
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                p = doc.add_heading(line[5:], level=4)
            elif line.startswith('##### '):
                p = doc.add_heading(line[6:], level=5)
            elif line.startswith('###### '):
                p = doc.add_heading(line[7:], level=6)

            # Lists
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif len(line) > 2 and line[0].isdigit() and line[1:3] == '. ':
                doc.add_paragraph(line[3:], style='List Number')

            # Blockquote
            elif line.startswith('> '):
                p = doc.add_paragraph(line[2:])
                p.style = 'Quote'

            # Empty line
            elif line.strip() == '':
                doc.add_paragraph()

            # Regular paragraph
            else:
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _add_formatted_text(self, paragraph, text: str):
        """Add text with inline formatting (bold, italic, code) to paragraph."""
        import re

        # Pattern for **bold**, *italic*, `code`, [link](url)
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[([^\]]+)\]\(([^\)]+)\))'

        parts = re.split(pattern, text)

        for i, part in enumerate(parts):
            if not part:
                continue

            if part.startswith('**') and part.endswith('**'):
                # Bold
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                # Italic
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                # Code
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            elif part.startswith('['):
                # This is handled by the regex groups
                continue
            elif i > 0 and parts[i-1] and parts[i-1].startswith('['):
                # Link text (captured group)
                run = paragraph.add_run(part)
                run.font.color.rgb = None  # Blue color would need RGBColor
            else:
                # Regular text
                paragraph.add_run(part)

    def export_to_pdf(self, markdown_text: str) -> bytes:
        """
        Export Markdown to PDF format.

        Args:
            markdown_text: Markdown content

        Returns:
            PDF file as bytes
        """
        # Convert markdown to HTML
        html_content = markdown.markdown(
            markdown_text,
            extensions=['extra', 'codehilite', 'tables', 'fenced_code']
        )

        # Wrap in HTML document with styling
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 2.5cm;
                }}
                body {{
                    font-family: 'DejaVu Sans', Arial, sans-serif;
                    font-size: 11pt;
                    line-height: 1.6;
                    color: #333;
                }}
                h1 {{
                    font-size: 24pt;
                    margin-top: 20pt;
                    margin-bottom: 12pt;
                    border-bottom: 2px solid #333;
                    padding-bottom: 6pt;
                }}
                h2 {{
                    font-size: 18pt;
                    margin-top: 16pt;
                    margin-bottom: 10pt;
                }}
                h3 {{
                    font-size: 14pt;
                    margin-top: 12pt;
                    margin-bottom: 8pt;
                }}
                p {{
                    margin-bottom: 10pt;
                }}
                code {{
                    background-color: #f4f4f4;
                    padding: 2pt 4pt;
                    border-radius: 3pt;
                    font-family: 'DejaVu Sans Mono', 'Courier New', monospace;
                    font-size: 10pt;
                }}
                pre {{
                    background-color: #f4f4f4;
                    padding: 12pt;
                    border-radius: 4pt;
                    overflow-x: auto;
                    margin: 10pt 0;
                }}
                pre code {{
                    background-color: transparent;
                    padding: 0;
                }}
                blockquote {{
                    border-left: 4pt solid #ddd;
                    padding-left: 12pt;
                    margin-left: 0;
                    font-style: italic;
                    color: #666;
                }}
                ul, ol {{
                    margin-bottom: 10pt;
                    padding-left: 20pt;
                }}
                li {{
                    margin-bottom: 4pt;
                }}
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 10pt 0;
                }}
                th, td {{
                    border: 1pt solid #ddd;
                    padding: 8pt;
                    text-align: left;
                }}
                th {{
                    background-color: #f4f4f4;
                    font-weight: bold;
                }}
                a {{
                    color: #0066cc;
                    text-decoration: none;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """

        # Convert HTML to PDF
        pdf_bytes = HTML(string=full_html).write_pdf()

        return pdf_bytes

    def export_to_rtf(self, markdown_text: str) -> bytes:
        """
        Export Markdown to RTF format.

        Args:
            markdown_text: Markdown content

        Returns:
            RTF file as bytes
        """
        # RTF header with UTF-8 support
        rtf = r'{\rtf1\ansi\ansicpg1252\deff0\nouicompat\deflang1033'
        rtf += r'{\fonttbl{\f0\fswiss\fcharset0 Arial;}{\f1\fmodern\fcharset0 Courier New;}}'
        rtf += r'{\colortbl ;\red0\green0\blue0;\red102\green102\blue102;}'
        rtf += '\n'

        lines = markdown_text.split('\n')

        for line in lines:
            line = line.rstrip()

            # Convert special characters to RTF unicode
            line = self._escape_rtf(line)

            # Headers
            if line.startswith('# '):
                rtf += r'\pard\sa200\sl276\slmult1\b\fs32 ' + line[2:] + r'\b0\fs22\par' + '\n'
            elif line.startswith('## '):
                rtf += r'\pard\sa200\sl276\slmult1\b\fs28 ' + line[3:] + r'\b0\fs22\par' + '\n'
            elif line.startswith('### '):
                rtf += r'\pard\sa200\sl276\slmult1\b\fs24 ' + line[4:] + r'\b0\fs22\par' + '\n'

            # Lists
            elif line.startswith('- ') or line.startswith('* '):
                rtf += r'\pard\fi-360\li720\sa200\sl276\slmult1 ' + r'\bullet\tab ' + line[2:] + r'\par' + '\n'

            # Blockquote
            elif line.startswith('> '):
                rtf += r'\pard\li720\sa200\sl276\slmult1\i ' + line[2:] + r'\i0\par' + '\n'

            # Empty line
            elif line.strip() == '':
                rtf += r'\par' + '\n'

            # Regular paragraph
            else:
                rtf += r'\pard\sa200\sl276\slmult1 ' + line + r'\par' + '\n'

        rtf += '}'

        return rtf.encode('utf-8')

    def _escape_rtf(self, text: str) -> str:
        """Escape special characters and convert Unicode to RTF format."""
        result = []

        for char in text:
            code = ord(char)

            # ASCII special characters
            if char == '\\':
                result.append('\\\\')
            elif char == '{':
                result.append('\\{')
            elif char == '}':
                result.append('\\}')
            # Unicode characters (including emojis)
            elif code > 127:
                result.append(f'\\u{code}?')
            else:
                result.append(char)

        return ''.join(result)


# Global converter instance
# Use environment variables for configuration
# Default to remote Ollama server for better performance
ollama_url = os.getenv('OLLAMA_URL', 'http://sonne.lan:8000')
ollama_model = os.getenv('OLLAMA_MODEL', 'gemma3:12b')
ollama_vision_model = os.getenv('OLLAMA_VISION_MODEL', 'gemma3:12b')

converter = MarkdownConverter(
    ollama_url=ollama_url,
    ollama_model=ollama_model,
    ollama_vision_model=ollama_vision_model
)
